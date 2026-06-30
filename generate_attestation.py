"""
Genere une Attestation de Reussite (.docx) par etudiant a partir du gabarit
AR.docx. Le gabarit n'a pas de champs de fusion : on repere les paragraphes
par leur libelle et on complete les runs vides (ou on remplace le texte d'un
run existant) en conservant la mise en forme d'origine.

Champs laisses en remplissage manuel (comme la decision du jury dans le
releve) : la mention ("Avec la mention :") et la date de signature
("TUNIS, le ...").
"""
from pathlib import Path

import docx


def _find_paragraph(doc, label):
    for p in doc.paragraphs:
        if label in p.text:
            return p
    return None


def _append_value(paragraph, value):
    """Ajoute un run avec value a la fin du paragraphe, en reprenant la mise
    en forme (police/taille/gras) du dernier run existant."""
    if paragraph is None or not value:
        return
    ref = paragraph.runs[-1] if paragraph.runs else None
    run = paragraph.add_run(str(value))
    if ref is not None:
        run.bold = ref.bold
        run.italic = ref.italic
        run.underline = ref.underline
        run.font.name = ref.font.name
        run.font.size = ref.font.size


def _replace_run_containing(paragraph, substring, new_text):
    """Remplace le texte du 1er run du paragraphe contenant substring."""
    if paragraph is None:
        return False
    for run in paragraph.runs:
        if substring in run.text:
            run.text = run.text.replace(substring, new_text)
            return True
    return False


def fill_attestation(template_path, output_path, nom_prenom, *, naissance_date=None,
                      cin="", nationalite="", cycle="", niveau_num="", niveau_suffixe="",
                      specialite=""):
    doc = docx.Document(template_path)

    _append_value(_find_paragraph(doc, "Nom et Prénom"), nom_prenom)

    naissance_str = naissance_date.strftime("%d/%m/%Y") if hasattr(naissance_date, "strftime") else (naissance_date or "")
    _append_value(_find_paragraph(doc, "Né (e) le") or _find_paragraph(doc, "Né  ( e )"), naissance_str)

    nationalite_p = _find_paragraph(doc, "Nationalité")
    if nationalite_p is not None:
        _replace_run_containing(nationalite_p, "Tunisienne", nationalite or "Tunisienne")

    _append_value(_find_paragraph(doc, "CIN"), cin)

    succes_p = _find_paragraph(doc, "A subi avec succès")
    if succes_p is not None and len(succes_p.runs) >= 4:
        runs = succes_p.runs
        runs[0].text = runs[0].text.rsplit(" ", 1)[0] + f" {niveau_num}" if niveau_num else runs[0].text
        runs[1].text = niveau_suffixe or runs[1].text
        runs[2].text = f" année {cycle} en " if cycle else runs[2].text
        runs[3].text = specialite or runs[3].text

    option_p = _find_paragraph(doc, "Option")
    if option_p is not None and len(option_p.runs) >= 2:
        option_p.runs[1].text = specialite or option_p.runs[1].text

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
